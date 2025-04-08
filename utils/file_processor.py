import streamlit as st
import pandas as pd
import numpy as np
import os
import json
import io
import docx
import re
from datetime import datetime
import tabula
import PyPDF2
import openpyxl
import xlrd

# Supported file types
supported_file_types = [
    ".csv", ".xlsx", ".xls", 
    ".json", ".txt", ".docx", 
    ".pdf"
]

def is_valid_file_type(filename):
    """Check if the uploaded file is of a supported type."""
    file_extension = os.path.splitext(filename)[1].lower()
    return file_extension in supported_file_types

def detect_file_type(filename):
    """Detect the file type based on the file extension."""
    file_extension = os.path.splitext(filename)[1].lower()
    return file_extension

def read_csv(file):
    """Read and parse CSV files."""
    try:
        return pd.read_csv(file)
    except Exception as e:
        # Try different encodings and delimiters if default fails
        try:
            return pd.read_csv(file, encoding='latin1')
        except:
            try:
                return pd.read_csv(file, sep=';')
            except:
                st.error(f"Failed to load CSV file: {str(e)}")
                return None

def read_excel(file):
    """Read and parse Excel files."""
    try:
        return pd.read_excel(file)
    except Exception as e:
        st.error(f"Failed to load Excel file: {str(e)}")
        return None

def read_json(file):
    """Read and parse JSON files."""
    try:
        # Try to parse as records (list of dictionaries)
        data = json.loads(file.read().decode('utf-8'))
        
        # Handle different JSON structures
        if isinstance(data, list):
            return pd.json_normalize(data)
        elif isinstance(data, dict):
            # Check if it's a nested structure
            if any(isinstance(v, (dict, list)) for v in data.values()):
                return pd.json_normalize(data)
            else:
                return pd.DataFrame([data])
        else:
            st.error("Unsupported JSON structure")
            return None
            
    except Exception as e:
        st.error(f"Failed to load JSON file: {str(e)}")
        return None

def read_text(file):
    """Read and parse text files, attempting to detect structure."""
    try:
        content = file.read().decode('utf-8')
        
        # Try to detect if it's a CSV
        if ',' in content.split('\n')[0]:
            return pd.read_csv(io.StringIO(content))
        
        # Try to detect if it's tab-separated
        elif '\t' in content.split('\n')[0]:
            return pd.read_csv(io.StringIO(content), sep='\t')
        
        # Try to detect if it's semicolon-separated
        elif ';' in content.split('\n')[0]:
            return pd.read_csv(io.StringIO(content), sep=';')
        
        # If no structure detected, create a DataFrame with text content
        else:
            lines = content.split('\n')
            return pd.DataFrame({'text': lines})
            
    except Exception as e:
        st.error(f"Failed to load text file: {str(e)}")
        return None

def read_docx(file):
    """Extract text and tables from Word documents."""
    try:
        # Save the uploaded file temporarily
        bytes_data = file.read()
        with open("temp.docx", "wb") as f:
            f.write(bytes_data)
        
        # Parse the document
        doc = docx.Document("temp.docx")
        
        # Extract text
        text_content = []
        for para in doc.paragraphs:
            text_content.append(para.text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                tables.append(pd.DataFrame(table_data[1:], columns=table_data[0]))
        
        # Clean up
        os.remove("temp.docx")
        
        # Return the first table if available, otherwise create a DataFrame from text
        if tables:
            return tables[0]
        else:
            return pd.DataFrame({'text': text_content})
            
    except Exception as e:
        st.error(f"Failed to load Word document: {str(e)}")
        if os.path.exists("temp.docx"):
            os.remove("temp.docx")
        return None

def read_pdf(file):
    """Extract text and tables from PDF documents."""
    try:
        # Save the uploaded file temporarily
        bytes_data = file.read()
        with open("temp.pdf", "wb") as f:
            f.write(bytes_data)
        
        # Extract tables using tabula
        tables = tabula.read_pdf("temp.pdf", pages='all')
        
        # Extract text using PyPDF2
        pdf_reader = PyPDF2.PdfReader("temp.pdf")
        text_content = []
        for page in pdf_reader.pages:
            text_content.append(page.extract_text())
        
        # Clean up
        os.remove("temp.pdf")
        
        # Return the first table if available, otherwise create a DataFrame from text
        if tables and not tables[0].empty:
            return tables[0]
        else:
            return pd.DataFrame({'text': text_content})
            
    except Exception as e:
        st.error(f"Failed to load PDF document: {str(e)}")
        if os.path.exists("temp.pdf"):
            os.remove("temp.pdf")
        return None

def process_uploaded_file(uploaded_file):
    """Process the uploaded file based on its type."""
    if uploaded_file is None:
        return None
        
    # Check file size limit based on subscription tier
    from utils.access_control import check_access
    
    # Get file size in MB
    file_size_mb = uploaded_file.size / (1024 * 1024)
    
    # Check if file size is within the user's limit
    if not check_access("file_size_limit", file_size_mb):
        st.error(f"File size ({file_size_mb:.2f} MB) exceeds your subscription tier limit. Please upgrade your plan to upload larger files.")
        return None
        
    file_type = detect_file_type(uploaded_file.name)
    
    # Process based on file type
    if file_type == '.csv':
        df = read_csv(uploaded_file)
    elif file_type in ['.xlsx', '.xls']:
        df = read_excel(uploaded_file)
    elif file_type == '.json':
        df = read_json(uploaded_file)
    elif file_type == '.txt':
        df = read_text(uploaded_file)
    elif file_type == '.docx':
        df = read_docx(uploaded_file)
    elif file_type == '.pdf':
        df = read_pdf(uploaded_file)
    else:
        st.error(f"Unsupported file type: {file_type}")
        return None
    
    # Validate the DataFrame
    if df is not None and len(df) > 0:
        return df
    else:
        st.error("No data could be extracted from the file.")
        return None

def infer_column_types(df):
    """Infer the data types of each column."""
    column_types = {}
    
    for column in df.columns:
        # Skip columns with all missing values
        if df[column].isna().all():
            column_types[column] = 'unknown'
            continue
        
        # Special handling for columns that might contain dates or deadlines in their name
        if any(keyword in column.lower() for keyword in ['date', 'deadline', 'time', 'day', 'month', 'year']):
            # First try to check if it looks like a date column
            sample = df[column].dropna().iloc[0] if not df[column].dropna().empty else None
            if sample and isinstance(sample, str):
                try:
                    pd.to_datetime(df[column])
                    column_types[column] = 'datetime'
                    continue
                except:
                    # Even if conversion fails, if it has "date" or "deadline" in name, treat as text
                    column_types[column] = 'text'
                    continue
            
        # Check if it's a numeric column
        try:
            pd.to_numeric(df[column])
            # Check if it's a binary/boolean column
            unique_values = df[column].dropna().unique()
            if len(unique_values) <= 2 and all(v in [0, 1, True, False] for v in unique_values):
                column_types[column] = 'boolean'
            else:
                column_types[column] = 'numeric'
            continue
        except:
            pass
            
        # Check if it's a datetime column for non-date-named columns
        try:
            pd.to_datetime(df[column])
            column_types[column] = 'datetime'
            continue
        except:
            pass
            
        # Check if it's a categorical column
        unique_count = df[column].nunique()
        if unique_count < 0.2 * len(df) and unique_count < 50:
            column_types[column] = 'categorical'
            continue
            
        # Default to text/string
        column_types[column] = 'text'
    
    return column_types

def apply_column_types(df, column_types):
    """Apply the detected or user-specified column types to the DataFrame."""
    df_copy = df.copy()
    
    for column, col_type in column_types.items():
        if column not in df_copy.columns:
            continue
            
        if col_type == 'numeric':
            try:
                df_copy[column] = pd.to_numeric(df_copy[column], errors='coerce')
            except:
                pass
        elif col_type == 'datetime':
            try:
                df_copy[column] = pd.to_datetime(df_copy[column], errors='coerce')
            except:
                pass
        elif col_type == 'boolean':
            try:
                # Convert various forms of True/False
                df_copy[column] = df_copy[column].map({
                    'True': True, 'False': False,
                    'true': True, 'false': False,
                    'TRUE': True, 'FALSE': False,
                    '1': True, '0': False,
                    1: True, 0: False,
                    'Yes': True, 'No': False,
                    'yes': True, 'no': False,
                    'Y': True, 'N': False,
                    'y': True, 'n': False
                }).astype('boolean')
            except:
                pass
    
    return df_copy
